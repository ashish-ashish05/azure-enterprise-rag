from io import BytesIO

from docx import Document as DocxDocument

from enterprise_rag.domain.models import Document
from enterprise_rag.ingestion.loaders.base import DocumentLoader


class DocxDocumentLoader(DocumentLoader):
    """Extract text from DOCX documents."""

    def load(
        self,
        content: bytes,
        *,
        document_id: str,
        document_family_id: str,
        source: str,
    ) -> Document:
        document = DocxDocument(BytesIO(content))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        full_text = "\n\n".join(paragraphs)

        return Document(
            document_id=document_id,
            document_family_id=document_family_id,
            source=source,
            content=full_text,
            content_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )